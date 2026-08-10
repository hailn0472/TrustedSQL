from __future__ import annotations

import argparse
import csv
import json
import math
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SETTING_LABELS = {
    "full_trustedsql": "Full TrustedSQL",
    "trustedsql_minus_m1": "Minus M1",
    "trustedsql_minus_m2": "Minus M2",
    "trustedsql_minus_m3_m4_m5": "Minus M3-M4-M5",
    "trustedsql_minus_m7": "Minus M7",
}

SETTING_LONG = {
    "full_trustedsql": "Full TrustedSQL",
    "trustedsql_minus_m1": "Minus M1 - Prompt Integrity Guard",
    "trustedsql_minus_m2": "Minus M2 - Conversation-Risk Model",
    "trustedsql_minus_m3_m4_m5": "Minus M3-M4-M5 - Authorization Block",
    "trustedsql_minus_m7": "Minus M7 - SQL Conformance Validator",
}

PALETTE = {
    "navy": (31, 58, 95),
    "blue": (46, 116, 181),
    "light_blue": (232, 238, 245),
    "gray": (98, 105, 112),
    "light_gray": (242, 244, 247),
    "green": (47, 111, 78),
    "red": (179, 58, 58),
    "gold": (196, 138, 44),
    "purple": (146, 112, 184),
    "teal": (77, 138, 158),
    "ink": (25, 30, 36),
    "white": (255, 255, 255),
}

METRIC_SHORT = {
    "RBAC Single-Turn Security - ASR \u2193": "RBAC ASR",
    "Prompt Injection Single-Turn Security - ASR \u2193": "PI ASR",
    "Multi-Turn Security - Sequence ASR \u2193": "Seq ASR",
    "Multi-Turn Security - Valid Secure Sequence Rate \u2191": "VSSR",
    "Utility - ST-EX \u2191": "ST-EX",
    "Utility - MT-IEX \u2191": "MT-IEX",
}


def main() -> int:
    parser = argparse.ArgumentParser(description="Build DOCX report for EX3 ablation dataflow analysis.")
    parser.add_argument("--analysis-dir", default="output/ex3_analysis")
    parser.add_argument("--output", default="output/docx/TrustedSQL_EX3_Ablation_Dataflow_Analysis.docx")
    parser.add_argument("--project-root", default=None)
    args = parser.parse_args()

    root = Path(args.project_root).resolve() if args.project_root else Path.cwd()
    analysis_dir = resolve(root, args.analysis_dir)
    output_docx = resolve(root, args.output)
    figures_dir = analysis_dir / "docx_figures"
    figures_dir.mkdir(parents=True, exist_ok=True)
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    data = load_analysis_data(analysis_dir)
    figures = build_figures(data, figures_dir)
    doc = build_document(data, figures)
    doc.save(output_docx)
    print(f"Wrote {output_docx}")
    return 0


def resolve(root: Path, value: str) -> Path:
    path = Path(value)
    return path.resolve() if path.is_absolute() else (root / path).resolve()


def load_analysis_data(analysis_dir: Path) -> dict[str, list[dict[str, Any]]]:
    names = [
        "aggregate_metric_summary.csv",
        "turn_decision_counts.csv",
        "blocked_stage_by_source.csv",
        "pairwise_decision_migrations.csv",
        "sequence_security_summary.csv",
        "utility_evidence_summary.csv",
        "run_coverage.csv",
        "module_reach_counts.csv",
        "pipeline_path_counts.csv",
        "sample_decision_matrix.csv",
    ]
    data: dict[str, list[dict[str, Any]]] = {}
    for name in names:
        data[name] = read_csv(analysis_dir / name)
    case_path = analysis_dir / "case_studies.json"
    data["case_studies.json"] = json.loads(case_path.read_text(encoding="utf-8")) if case_path.exists() else []
    return data


def build_figures(data: dict[str, list[dict[str, Any]]], figures_dir: Path) -> dict[str, Path]:
    figures = {
        "aggregate_delta": figures_dir / "figure_1_security_delta_vs_full.png",
        "decision_distribution": figures_dir / "figure_2_attack_decision_distribution.png",
        "blocking_stage": figures_dir / "figure_3_blocking_stage_distribution.png",
        "migration": figures_dir / "figure_4_pairwise_migration_highlights.png",
        "sequence_security": figures_dir / "figure_5_sequence_security.png",
        "module_reach": figures_dir / "figure_6_module_reach_heatmap.png",
    }
    draw_aggregate_delta(data["aggregate_metric_summary.csv"], figures["aggregate_delta"])
    draw_decision_distribution(data["turn_decision_counts.csv"], figures["decision_distribution"])
    draw_blocking_stage(data["blocked_stage_by_source.csv"], figures["blocking_stage"])
    draw_migration_highlights(data["pairwise_decision_migrations.csv"], figures["migration"])
    draw_sequence_security(data["sequence_security_summary.csv"], figures["sequence_security"])
    draw_module_reach_heatmap(data["module_reach_counts.csv"], figures["module_reach"])
    return figures


def build_document(data: dict[str, list[dict[str, Any]]], figures: dict[str, Path]) -> Document:
    doc = Document()
    configure_document(doc)
    add_masthead(doc)
    add_executive_summary(doc, data)
    add_section_aggregate(doc, data, figures["aggregate_delta"])
    add_section_runtime(doc, data, figures["decision_distribution"], figures["blocking_stage"])
    add_section_internal_framework(doc, data, figures["module_reach"])
    add_section_case_studies(doc, data)
    add_section_migration(doc, data, figures["migration"])
    add_section_sequence(doc, data, figures["sequence_security"])
    add_section_paper_text(doc)
    add_appendix(doc, data)
    return doc


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    caption = styles.add_style("Figure Caption", 1)
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(80, 80, 80)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)

    small = styles.add_style("Small Note", 1)
    small.font.name = "Calibri"
    small._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    small._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    small.font.size = Pt(9.5)
    small.font.color.rgb = RGBColor(80, 80, 80)
    small.paragraph_format.space_after = Pt(4)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("TrustedSQL EX3 Ablation Analysis")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(98, 105, 112)


def add_masthead(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("TRUSTEDSQL RESULT ANALYSIS")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(98, 105, 112)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("EX3 Ablation Dataflow and Security Impact")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 58, 95)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(
        "A focused results memo explaining how TrustedSQL request routing, blocking stages, and sequence-level security change when framework components are removed."
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)

    rows = [
        ("Scope", "EX3 component ablation, 3 repeated Gemini 2.5 Flash runs"),
        ("Inputs", "output/result_all_3times.csv and output/ex3 raw runtime traces"),
        ("Output", "Figures, tables, and paper-ready interpretation for Results/Discussion"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [1800, 7560])
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        shade_cell(cells[0], "F2F4F7")
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    mark_header_row(table.rows[0])
    add_bottom_rule(doc)


def add_executive_summary(doc: Document, data: dict[str, list[dict[str, Any]]]) -> None:
    doc.add_heading("1. Executive Summary", level=1)
    add_callout(
        doc,
        "Core finding",
        "The EX3 ablation results should be framed as a dataflow story rather than a simple metric table. Removing each component changes where requests stop, which requests continue to SQL generation/execution, and which security objective loses protection.",
    )
    add_bullet(
        doc,
        "M1 is the primary prompt-injection entry gate: removing it increases Prompt Injection ASR from 1.57% to 31.13%.",
    )
    add_bullet(
        doc,
        "M2 controls conversation-aware risk: removing it improves utility but weakens sequence-level security, lowering VSSR by 3.57 percentage points.",
    )
    add_bullet(
        doc,
        "M3-M4-M5 is the deterministic authorization block: removing it is the strongest failure mode, increasing RBAC ASR by 58.25 percentage points.",
    )
    add_bullet(
        doc,
        "M7 behaves as a final SQL-conformance backstop: removing it has the smallest impact on RBAC and prompt-injection metrics.",
    )
    p = doc.add_paragraph(style="Small Note")
    p.add_run("Important caveat: ").bold = True
    p.add_run(
        "the current output/ex3 folder contains raw traces for the four ablation configurations, but not the raw Full TrustedSQL run. Therefore, full-vs-ablation comparisons use aggregate metrics, while per-request migrations are computed among the raw ablation runs that are present."
    )


def add_section_aggregate(doc: Document, data: dict[str, list[dict[str, Any]]], figure: Path) -> None:
    doc.add_heading("2. Aggregate Security and Utility Deltas", level=1)
    doc.add_paragraph(
        "The aggregate table shows how each ablation moves the headline metrics relative to Full TrustedSQL. This view is useful for paper-level claims, but it does not by itself explain why the numbers move."
    )
    add_picture(doc, figure, "Figure 1. Aggregate delta against Full TrustedSQL for selected EX3 metrics.")
    rows = data["aggregate_metric_summary.csv"]
    table_rows = []
    for setting in ["trustedsql_minus_m1", "trustedsql_minus_m2", "trustedsql_minus_m3_m4_m5", "trustedsql_minus_m7"]:
        table_rows.append(
            [
                SETTING_LONG[setting],
                metric_value(rows, setting, "RBAC Single-Turn Security - ASR \u2193"),
                metric_delta(rows, setting, "RBAC Single-Turn Security - ASR \u2193"),
                metric_value(rows, setting, "Prompt Injection Single-Turn Security - ASR \u2193"),
                metric_delta(rows, setting, "Prompt Injection Single-Turn Security - ASR \u2193"),
                metric_value(rows, setting, "Multi-Turn Security - Valid Secure Sequence Rate \u2191"),
                metric_delta(rows, setting, "Multi-Turn Security - Valid Secure Sequence Rate \u2191"),
            ]
        )
    add_table(
        doc,
        ["Configuration", "RBAC ASR", "Delta", "PI ASR", "Delta", "VSSR", "Delta"],
        table_rows,
        widths=[2700, 1050, 900, 1050, 900, 1050, 1710],
    )
    doc.add_paragraph(
        "The largest movement is the authorization-block ablation. Its RBAC ASR rises to 68.75%, which means that many requests that should be denied are no longer stopped before generation or execution. By contrast, the M7 ablation changes most security metrics only slightly, indicating that earlier gates carry most of the security load."
    )


def add_section_runtime(doc: Document, data: dict[str, list[dict[str, Any]]], decision_fig: Path, blocked_fig: Path) -> None:
    doc.add_heading("3. Runtime Dataflow: Decisions and Blocking Stages", level=1)
    doc.add_paragraph(
        "The raw runtime traces expose the operational mechanism behind the aggregate metrics. Each turn records its final decision, whether it executed, and the module where it stopped. This allows us to describe the framework as a sequence of gates rather than as a monolithic classifier."
    )
    add_picture(doc, decision_fig, "Figure 2. Decision distribution over attack-related EX3 turns, aggregated across the three raw ablation runs.")
    add_picture(doc, blocked_fig, "Figure 3. Blocking-stage redistribution across ablations. NONE means the request was not blocked before execution path completion.")

    blocked = data["blocked_stage_by_source.csv"]
    rows = [
        [
            "Minus M1",
            "ST-PI",
            top_blocked(blocked, "trustedsql_minus_m1", "ST-PI"),
            "Without the prompt-integrity gate, many prompt-injection attempts move downstream and are caught later by M4/M5 or reach ALLOW.",
        ],
        [
            "Minus M2",
            "MT-MAL",
            top_blocked(blocked, "trustedsql_minus_m2", "MT-MAL"),
            "Conversation-aware denial is reduced; downstream policy checks absorb part of the risk, but sequence ASR still rises.",
        ],
        [
            "Minus M3-M4-M5",
            "ST-RBAC",
            top_blocked(blocked, "trustedsql_minus_m3_m4_m5", "ST-RBAC"),
            "The deterministic authorization gates disappear, so RBAC attacks are no longer concentrated at M4/M5 and many proceed.",
        ],
        [
            "Minus M7",
            "All attack groups",
            "Mostly M1/M5/M2; few stops at execution",
            "The main security decisions remain upstream; M7 mainly catches SQL-level conformance problems.",
        ],
    ]
    add_table(
        doc,
        ["Ablation", "Most relevant traffic", "Dominant stops", "Interpretation"],
        rows,
        widths=[1350, 1450, 2300, 4260],
    )


def add_section_internal_framework(doc: Document, data: dict[str, list[dict[str, Any]]], module_fig: Path) -> None:
    doc.add_heading("4. Internal Framework Data Representation", level=1)
    doc.add_paragraph(
        "The deeper trace view shows that an ablation changes not only the final decision, but also the internal representation that is available to downstream modules. When a gate is removed, later modules receive a different context: missing risk evidence, missing resource contracts, missing authorization proofs, or missing SQL conformance checks."
    )
    add_picture(doc, module_fig, "Figure 4. Module reach heatmap for attack-related turns. Darker cells mean more turns reached that module across the three raw EX3 runs.")

    path_rows = data["pipeline_path_counts.csv"]
    rows = []
    for setting in ["trustedsql_minus_m1", "trustedsql_minus_m2", "trustedsql_minus_m3_m4_m5", "trustedsql_minus_m7"]:
        top_paths = [
            row
            for row in path_rows
            if row["setting_id"] == setting and row["source_group"] in {"ST-RBAC", "ST-PI", "MT-MAL"}
        ]
        top_paths.sort(key=lambda row: int(float(row["turns"])), reverse=True)
        compact_paths = "; ".join(
            f"{row['source_group']} {row['decision']}@{row['blocked_at']} ({int(float(row['turns']))})"
            for row in top_paths[:3]
        )
        rows.append([SETTING_LONG[setting], compact_paths, representation_change(setting)])
    add_table(
        doc,
        ["Configuration", "Dominant raw paths", "What changes inside the framework"],
        rows,
        widths=[2400, 3300, 3660],
    )


def add_section_case_studies(doc: Document, data: dict[str, list[dict[str, Any]]]) -> None:
    doc.add_heading("5. Representative Per-Request Case Studies", level=1)
    doc.add_paragraph(
        "The following examples trace the same request through two ablation configurations. They are useful for explaining causal mechanism: which artifact was produced, which gate stopped the request, and what happened when that gate was absent."
    )
    cases = data.get("case_studies.json", [])
    for case in cases:
        doc.add_heading(f"{case['case_id']}. {case['title']}", level=2)
        add_callout(doc, "Request", f"{case['sample_id']} turn {case['turn_id']}: {case['nlq']}")
        rows = [
            [
                SETTING_LONG.get(case["left"]["setting_id"], case["left"]["setting_id"]),
                case["left"]["decision"],
                case["left"]["blocked_at"],
                case["left"]["trace_path"],
            ],
            [
                SETTING_LONG.get(case["right"]["setting_id"], case["right"]["setting_id"]),
                case["right"]["decision"],
                case["right"]["blocked_at"],
                case["right"]["trace_path"],
            ],
        ]
        add_table(doc, ["Configuration", "Decision", "Blocked at", "Trace path"], rows, widths=[2300, 1000, 1100, 4960])

        evidence_rows = []
        for side_name in ["left", "right"]:
            side = case[side_name]
            setting = SETTING_LABELS.get(side["setting_id"], side["setting_id"])
            evidence_rows.extend(case_module_evidence_rows(setting, side["module_summaries"]))
        add_table(
            doc,
            ["Setting", "Module", "Evidence extracted from artifact/audit"],
            evidence_rows[:10],
            widths=[1600, 900, 6860],
        )
        if case["right"].get("final_sql"):
            p = doc.add_paragraph(style="Small Note")
            p.add_run("Executed SQL excerpt: ").bold = True
            p.add_run(case["right"]["final_sql"])
        doc.add_paragraph(case["interpretation"])


def add_section_migration(doc: Document, data: dict[str, list[dict[str, Any]]], figure: Path) -> None:
    doc.add_heading("6. Per-Request Migration Evidence", level=1)
    doc.add_paragraph(
        "Because the raw Full TrustedSQL traces are not present, direct Full-vs-ablation migration cannot yet be computed. However, pairwise comparisons among raw ablation runs still reveal how removing one block changes the fate of the same sample_id and turn_id."
    )
    add_picture(doc, figure, "Figure 5. Largest pairwise decision migrations among raw EX3 ablation traces.")
    rows = data["pairwise_decision_migrations.csv"]
    highlights = [
        ("trustedsql_minus_m2", "trustedsql_minus_m3_m4_m5", "ST-RBAC", "DENY", "ALLOW"),
        ("trustedsql_minus_m2", "trustedsql_minus_m3_m4_m5", "MT-MAL", "DENY", "ALLOW"),
        ("trustedsql_minus_m1", "trustedsql_minus_m2", "ST-PI", "ALLOW", "DENY"),
        ("trustedsql_minus_m3_m4_m5", "trustedsql_minus_m7", "ST-RBAC", "ALLOW", "DENY"),
    ]
    table_rows = []
    for left, right, source, left_decision, right_decision in highlights:
        turns = migration_turns(rows, left, right, source, left_decision, right_decision)
        table_rows.append(
            [
                f"{SETTING_LABELS[left]} -> {SETTING_LABELS[right]}",
                source,
                f"{left_decision} -> {right_decision}",
                str(turns),
                migration_interpretation(left, right, source, left_decision, right_decision),
            ]
        )
    add_table(
        doc,
        ["Comparison", "Traffic", "Decision shift", "Turns", "Meaning"],
        table_rows,
        widths=[2300, 1100, 1350, 800, 3810],
    )
    doc.add_paragraph(
        "The strongest migration is the authorization comparison: when the M3-M4-M5 block is absent, hundreds of ST-RBAC turns that were denied in the configuration with policy checks become allowed. This is stronger evidence than the aggregate ASR alone because it shows the same request crossing a different gate."
    )


def add_section_sequence(doc: Document, data: dict[str, list[dict[str, Any]]], figure: Path) -> None:
    doc.add_heading("7. Multi-Turn Security and Utility Trade-off", level=1)
    add_picture(doc, figure, "Figure 6. Multi-turn prefix utility, sequence ASR, and valid secure sequence rate by ablation.")
    seq = data["sequence_security_summary.csv"]
    rows = []
    for setting in ["trustedsql_minus_m1", "trustedsql_minus_m2", "trustedsql_minus_m3_m4_m5", "trustedsql_minus_m7"]:
        row = find_row(seq, setting_id=setting)
        rows.append(
            [
                SETTING_LONG[setting],
                fmt_pct(row["prefix_rs_rate"]),
                fmt_pct(row["sequence_asr"]),
                fmt_pct(row["valid_secure_sequence_rate"]),
                str(int(float(row["final_allow"]))),
            ]
        )
    add_table(
        doc,
        ["Configuration", "Prefix-RS", "Sequence ASR", "VSSR", "Final ALLOW"],
        rows,
        widths=[3000, 1400, 1500, 1300, 2160],
    )
    doc.add_paragraph(
        "The multi-turn result should be interpreted jointly. Removing M3-M4-M5 improves Prefix-RS because fewer benign-prefix turns are blocked, but the same ablation increases Sequence ASR to 34.05%. In other words, the pipeline becomes more permissive and apparently more useful on prefixes, while becoming less secure at the decisive malicious turn."
    )


def add_section_paper_text(doc: Document) -> None:
    doc.add_heading("8. Suggested Results/Discussion Text", level=1)
    paragraphs = [
        (
            "Instead of interpreting EX3 only as a component-level score table, the ablation should be read as a pipeline dataflow experiment. Each removed module changes the point at which requests exit the pipeline, which in turn changes the balance among benign utility, prompt-injection refusal, RBAC refusal, and multi-turn sequence safety."
        ),
        (
            "The Prompt Integrity Guard ablation shows the clearest prompt-injection effect. When M1 is removed, prompt-injection ASR increases sharply while benign utility remains unchanged. This indicates that M1 primarily acts as an early prompt-integrity gate rather than as a general utility component."
        ),
        (
            "The Conversation-Risk Model ablation produces a different pattern. Removing M2 increases single-turn and multi-turn utility, but this gain is accompanied by weaker sequence-level protection. The result supports the interpretation that conversation-risk modeling is most important when the current request must be judged against accumulated conversational context."
        ),
        (
            "The M3-M4-M5 ablation is the strongest evidence for the necessity of policy-grounded authorization. Removing resource planning, table/column validation, and row-scope verification causes RBAC ASR to rise dramatically. Runtime traces confirm the mechanism: requests that would otherwise stop at deterministic authorization checks continue toward SQL generation or execution."
        ),
        (
            "Finally, the M7 ablation has the smallest effect on headline security metrics. This does not make SQL conformance validation unnecessary; rather, it indicates that M7 behaves as a downstream backstop after policy and risk gates have already handled most security-relevant decisions."
        ),
    ]
    for text in paragraphs:
        doc.add_paragraph(text)


def add_appendix(doc: Document, data: dict[str, list[dict[str, Any]]]) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix A. Generated Analysis Files", level=1)
    files = [
        ("aggregate_metric_summary.csv", "Key EX3 metrics and deltas against Full TrustedSQL."),
        ("turn_decision_counts.csv", "ALLOW/DENY/ERROR counts by ablation and scenario family."),
        ("blocked_stage_by_source.csv", "Module where requests stop, by ablation and source group."),
        ("pairwise_decision_migrations.csv", "Per-request decision changes among available raw ablation runs."),
        ("sample_decision_matrix.csv", "Side-by-side decision, block stage, and trace path for each raw ablation turn."),
        ("pipeline_path_counts.csv", "Dominant end-to-end pipeline paths grouped by ablation and scenario family."),
        ("module_reach_counts.csv", "Counts showing how many attack-related turns reached each module."),
        ("case_studies.json", "Representative request-level traces with module artifact summaries."),
        ("sequence_security_summary.csv", "Prefix-RS, Sequence ASR, and VSSR from sequence evidence."),
        ("utility_evidence_summary.csv", "Turn-level utility evidence for benign and multi-turn records."),
    ]
    add_table(doc, ["File", "Purpose"], files, widths=[3000, 6360])
    doc.add_paragraph(
        "The analysis is reproducible with: python tools/analysis/ex3_ablation_dataflow.py followed by python tools/analysis/build_ex3_ablation_docx.py.",
        style="Small Note",
    )


def draw_aggregate_delta(rows: list[dict[str, Any]], path: Path) -> None:
    metrics = [
        "RBAC Single-Turn Security - ASR \u2193",
        "Prompt Injection Single-Turn Security - ASR \u2193",
        "Multi-Turn Security - Sequence ASR \u2193",
        "Multi-Turn Security - Valid Secure Sequence Rate \u2191",
    ]
    settings = ["trustedsql_minus_m1", "trustedsql_minus_m2", "trustedsql_minus_m3_m4_m5", "trustedsql_minus_m7"]
    series = {
        metric: [float(find_row(rows, setting_id=s, metric=metric)["delta_vs_full"]) for s in settings]
        for metric in metrics
    }
    draw_delta_dotplot(
        path,
        title="Aggregate metric movement vs Full TrustedSQL",
        groups=[SETTING_LABELS[s] for s in settings],
        series={METRIC_SHORT[k]: v for k, v in series.items()},
    )


def draw_decision_distribution(rows: list[dict[str, Any]], path: Path) -> None:
    settings = ["trustedsql_minus_m1", "trustedsql_minus_m2", "trustedsql_minus_m3_m4_m5", "trustedsql_minus_m7"]
    decisions = ["ALLOW", "DENY", "ERROR"]
    data = {s: Counter() for s in settings}
    for row in rows:
        if row["source_group"] in {"ST-RBAC", "ST-PI", "MT-MAL"}:
            data[row["setting_id"]][row["decision"]] += int(float(row["turns"]))
    draw_bubble_matrix(
        path,
        title="Attack-traffic decision distribution",
        groups=[SETTING_LABELS[s] for s in settings],
        columns=decisions,
        values=[[data[s][d] for d in decisions] for s in settings],
        colors=[PALETTE["green"], PALETTE["red"], PALETTE["purple"]],
        note="Bubble area encodes turn count across ST-RBAC, ST-PI, and MT-MAL.",
    )


def draw_blocking_stage(rows: list[dict[str, Any]], path: Path) -> None:
    settings = ["trustedsql_minus_m1", "trustedsql_minus_m2", "trustedsql_minus_m3_m4_m5", "trustedsql_minus_m7"]
    stages = ["NONE", "M1", "M2", "M3", "M4", "M5", "M6", "M7", "X1"]
    data = {s: Counter() for s in settings}
    for row in rows:
        if row["source_group"] in {"ST-RBAC", "ST-PI", "MT-MAL"}:
            data[row["setting_id"]][row["blocked_at"]] += int(float(row["turns"]))
    draw_stage_heatmap(
        path,
        title="Where attack-related turns stop in the pipeline",
        groups=[SETTING_LABELS[s] for s in settings],
        columns=stages,
        values=[[data[s][stage] for stage in stages] for s in settings],
    )


def draw_migration_highlights(rows: list[dict[str, Any]], path: Path) -> None:
    items = [
        ("M2 -> M3-M4-M5", "ST-RBAC DENY -> ALLOW", migration_turns(rows, "trustedsql_minus_m2", "trustedsql_minus_m3_m4_m5", "ST-RBAC", "DENY", "ALLOW")),
        ("M2 -> M3-M4-M5", "MT-MAL DENY -> ALLOW", migration_turns(rows, "trustedsql_minus_m2", "trustedsql_minus_m3_m4_m5", "MT-MAL", "DENY", "ALLOW")),
        ("M1 -> M2", "ST-PI ALLOW -> DENY", migration_turns(rows, "trustedsql_minus_m1", "trustedsql_minus_m2", "ST-PI", "ALLOW", "DENY")),
        ("M3-M4-M5 -> M7", "ST-RBAC ALLOW -> DENY", migration_turns(rows, "trustedsql_minus_m3_m4_m5", "trustedsql_minus_m7", "ST-RBAC", "ALLOW", "DENY")),
        ("M3-M4-M5 -> M7", "MT-MAL ALLOW -> DENY", migration_turns(rows, "trustedsql_minus_m3_m4_m5", "trustedsql_minus_m7", "MT-MAL", "ALLOW", "DENY")),
    ]
    draw_migration_tilemap(
        path,
        title="Largest pairwise decision migrations among raw ablation traces",
        items=[(f"{a}: {b}", v) for a, b, v in items],
    )


def draw_sequence_security(rows: list[dict[str, Any]], path: Path) -> None:
    settings = ["trustedsql_minus_m1", "trustedsql_minus_m2", "trustedsql_minus_m3_m4_m5", "trustedsql_minus_m7"]
    draw_sequence_slopegraph(
        path,
        title="Multi-turn sequence security evidence",
        groups=[SETTING_LABELS[s] for s in settings],
        series={
            "Prefix-RS": [float(find_row(rows, setting_id=s)["prefix_rs_rate"]) for s in settings],
            "Sequence ASR": [float(find_row(rows, setting_id=s)["sequence_asr"]) for s in settings],
            "VSSR": [float(find_row(rows, setting_id=s)["valid_secure_sequence_rate"]) for s in settings],
        },
    )


def draw_module_reach_heatmap(rows: list[dict[str, Any]], path: Path) -> None:
    settings = ["trustedsql_minus_m1", "trustedsql_minus_m2", "trustedsql_minus_m3_m4_m5", "trustedsql_minus_m7"]
    modules = ["C0", "M1", "M2", "M3", "M4", "M5", "M6", "M7", "X1"]
    values: dict[tuple[str, str], int] = defaultdict(int)
    for row in rows:
        if row["source_group"] == "ATTACK_ALL":
            values[(row["setting_id"], row["module_id"])] += int(float(row["turns_reached"]))
    max_value = max(values.values()) if values else 1
    width, height = 1400, 760
    image = Image.new("RGB", (width, height), PALETTE["white"])
    draw = ImageDraw.Draw(image)
    fonts = chart_fonts()
    draw.text((70, 28), "Module reach across attack-related turns", font=fonts["title"], fill=PALETTE["navy"])
    left, top = 300, 120
    cell_w, cell_h = 105, 74
    for j, module in enumerate(modules):
        draw.text((left + j * cell_w + 28, top - 34), module, font=fonts["axis"], fill=PALETTE["ink"])
    for i, setting in enumerate(settings):
        y = top + i * cell_h
        draw.text((60, y + 24), SETTING_LABELS[setting], font=fonts["axis"], fill=PALETTE["ink"])
        for j, module in enumerate(modules):
            x = left + j * cell_w
            value = values[(setting, module)]
            intensity = value / max_value
            fill = blend((239, 244, 248), PALETTE["blue"], intensity)
            draw.rectangle((x, y, x + cell_w - 6, y + cell_h - 8), fill=fill, outline=(220, 225, 232))
            text_color = PALETTE["white"] if intensity > 0.55 else PALETTE["ink"]
            label = str(value) if value else "-"
            bbox = draw.textbbox((0, 0), label, font=fonts["small"])
            draw.text(
                (x + (cell_w - 6 - (bbox[2] - bbox[0])) / 2, y + 23),
                label,
                font=fonts["small"],
                fill=text_color,
            )
    draw.text(
        (70, height - 55),
        "Counts aggregate ST-RBAC, ST-PI, and MT-MAL turns across the three raw EX3 ablation runs.",
        font=fonts["axis"],
        fill=PALETTE["gray"],
    )
    image.save(path)


def blend(a: tuple[int, int, int], b: tuple[int, int, int], t: float) -> tuple[int, int, int]:
    t = max(0.0, min(1.0, t))
    return tuple(int(a[i] + (b[i] - a[i]) * t) for i in range(3))


def draw_delta_dotplot(
    path: Path,
    *,
    title: str,
    groups: list[str],
    series: dict[str, list[float]],
) -> None:
    width, height = 1400, 860
    image = Image.new("RGB", (width, height), PALETTE["white"])
    draw = ImageDraw.Draw(image)
    fonts = chart_fonts()
    draw.text((70, 28), title, font=fonts["title"], fill=PALETTE["navy"])
    all_values = [value for values in series.values() for value in values]
    min_v = min(min(all_values), -10)
    max_v = max(max(all_values), 10)
    left, right, top, row_gap = 285, 1285, 125, 130
    axis_y = top + len(groups) * row_gap + 30
    colors = [PALETTE["red"], PALETTE["gold"], PALETTE["purple"], PALETTE["green"], PALETTE["teal"]]

    for tick in nice_ticks(min_v, max_v, 6):
        x = value_to_x(tick, min_v, max_v, left, right - left)
        draw.line((x, top - 15, x, axis_y - 34), fill=(226, 230, 235), width=1)
        draw.text((x - 18, axis_y - 15), f"{tick:g}", font=fonts["small"], fill=PALETTE["gray"])
    zero_x = value_to_x(0, min_v, max_v, left, right - left)
    draw.line((zero_x, top - 20, zero_x, axis_y - 30), fill=(145, 150, 158), width=2)
    draw.text((left, axis_y + 15), "percentage-point delta vs Full TrustedSQL", font=fonts["axis"], fill=PALETTE["gray"])

    series_items = list(series.items())
    for gi, group in enumerate(groups):
        y = top + gi * row_gap
        draw.text((65, y + 22), group, font=fonts["axis"], fill=PALETTE["ink"])
        draw.line((left, y + 34, right, y + 34), fill=(238, 241, 245), width=1)
        for si, (name, values) in enumerate(series_items):
            value = values[gi]
            x = value_to_x(value, min_v, max_v, left, right - left)
            offset = (si - (len(series_items) - 1) / 2) * 16
            yy = y + 34 + offset
            draw.line((zero_x, yy, x, yy), fill=colors[si], width=2)
            draw.ellipse((x - 7, yy - 7, x + 7, yy + 7), fill=colors[si], outline=PALETTE["white"], width=2)
            if abs(value) >= 20 or name in {"RBAC ASR", "VSSR"}:
                draw.text((x + 10, yy - 9), f"{value:+.1f}", font=fonts["small"], fill=PALETTE["ink"])

    legend_x, legend_y = left, height - 70
    for si, (name, _values) in enumerate(series_items):
        draw.ellipse((legend_x, legend_y, legend_x + 14, legend_y + 14), fill=colors[si])
        draw.text((legend_x + 22, legend_y - 3), name, font=fonts["axis"], fill=PALETTE["ink"])
        legend_x += 190
    image.save(path)


def draw_bubble_matrix(
    path: Path,
    *,
    title: str,
    groups: list[str],
    columns: list[str],
    values: list[list[int]],
    colors: list[tuple[int, int, int]],
    note: str,
) -> None:
    width, height = 1400, 700
    image = Image.new("RGB", (width, height), PALETTE["white"])
    draw = ImageDraw.Draw(image)
    fonts = chart_fonts()
    draw.text((70, 28), title, font=fonts["title"], fill=PALETTE["navy"])
    left, top = 360, 155
    col_gap, row_gap = 245, 96
    max_value = max(max(row) for row in values) or 1
    for j, column in enumerate(columns):
        x = left + j * col_gap
        draw.text((x - 28, top - 80), column, font=fonts["axis"], fill=PALETTE["ink"])
    for i, group in enumerate(groups):
        y = top + i * row_gap
        draw.text((70, y - 10), group, font=fonts["axis"], fill=PALETTE["ink"])
        for j, column in enumerate(columns):
            value = values[i][j]
            x = left + j * col_gap
            radius = 10 + 36 * math.sqrt(value / max_value)
            draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill=colors[j], outline=PALETTE["white"], width=2)
            label = str(value)
            bbox = draw.textbbox((0, 0), label, font=fonts["small"])
            draw.text((x - (bbox[2] - bbox[0]) / 2, y - 8), label, font=fonts["small"], fill=PALETTE["white"])
    draw.text((70, height - 70), note, font=fonts["axis"], fill=PALETTE["gray"])
    image.save(path)


def draw_stage_heatmap(
    path: Path,
    *,
    title: str,
    groups: list[str],
    columns: list[str],
    values: list[list[int]],
) -> None:
    width, height = 1400, 780
    image = Image.new("RGB", (width, height), PALETTE["white"])
    draw = ImageDraw.Draw(image)
    fonts = chart_fonts()
    draw.text((70, 28), title, font=fonts["title"], fill=PALETTE["navy"])
    left, top = 310, 125
    cell_w, cell_h = 108, 78
    max_value = max(max(row) for row in values) or 1
    for j, column in enumerate(columns):
        x = left + j * cell_w
        draw.text((x + 25, top - 38), column, font=fonts["axis"], fill=PALETTE["ink"])
    for i, group in enumerate(groups):
        y = top + i * cell_h
        draw.text((60, y + 24), group, font=fonts["axis"], fill=PALETTE["ink"])
        for j, column in enumerate(columns):
            x = left + j * cell_w
            value = values[i][j]
            if column == "NONE":
                fill = blend((238, 247, 242), PALETTE["green"], value / max_value)
            else:
                fill = blend((248, 241, 241), PALETTE["red"], value / max_value)
            draw.rectangle((x, y, x + cell_w - 7, y + cell_h - 9), fill=fill, outline=(220, 225, 232))
            if value:
                label = str(value)
                bbox = draw.textbbox((0, 0), label, font=fonts["small"])
                intensity = value / max_value
                text_color = PALETTE["white"] if intensity > 0.55 else PALETTE["ink"]
                draw.text((x + (cell_w - 7 - (bbox[2] - bbox[0])) / 2, y + 25), label, font=fonts["small"], fill=text_color)
            else:
                draw.text((x + 43, y + 25), "-", font=fonts["small"], fill=PALETTE["gray"])
    draw.text((70, height - 70), "Green NONE cells reached the end of the blocking stage; red cells indicate where requests stopped.", font=fonts["axis"], fill=PALETTE["gray"])
    image.save(path)


def draw_migration_tilemap(
    path: Path,
    *,
    title: str,
    items: list[tuple[str, int]],
) -> None:
    width, height = 1400, 760
    image = Image.new("RGB", (width, height), PALETTE["white"])
    draw = ImageDraw.Draw(image)
    fonts = chart_fonts()
    draw.text((70, 28), title, font=fonts["title"], fill=PALETTE["navy"])
    max_value = max(v for _label, v in items) or 1
    left, top = 90, 120
    tile_w, tile_h = 390, 150
    for idx, (label, value) in enumerate(items):
        col = idx % 3
        row = idx // 3
        x = left + col * (tile_w + 35)
        y = top + row * (tile_h + 42)
        intensity = value / max_value
        fill = blend((239, 244, 248), PALETTE["blue"], intensity)
        draw.rounded_rectangle((x, y, x + tile_w, y + tile_h), radius=0, fill=fill, outline=(210, 218, 228), width=2)
        value_text = str(value)
        draw.text((x + 18, y + 18), value_text, font=fonts["title"], fill=PALETTE["white"] if intensity > 0.55 else PALETTE["navy"])
        for line_i, line in enumerate(wrap_text(draw, label, fonts["axis"], tile_w - 36)[:3]):
            draw.text((x + 18, y + 70 + line_i * 24), line, font=fonts["axis"], fill=PALETTE["white"] if intensity > 0.55 else PALETTE["ink"])
    draw.text((70, height - 65), "Each tile is a same-sample decision shift observed among available raw ablation traces.", font=fonts["axis"], fill=PALETTE["gray"])
    image.save(path)


def draw_sequence_slopegraph(
    path: Path,
    *,
    title: str,
    groups: list[str],
    series: dict[str, list[float]],
) -> None:
    width, height = 1400, 780
    image = Image.new("RGB", (width, height), PALETTE["white"])
    draw = ImageDraw.Draw(image)
    fonts = chart_fonts()
    draw.text((70, 28), title, font=fonts["title"], fill=PALETTE["navy"])
    left, right, top, bottom = 190, 1260, 120, 620
    x_gap = (right - left) / (len(groups) - 1)
    colors = [PALETTE["teal"], PALETTE["red"], PALETTE["green"]]
    for tick in [0, 20, 40, 60, 80, 100]:
        y = value_to_y(tick, 0, 100, top, bottom - top)
        draw.line((left, y, right, y), fill=(232, 236, 241), width=1)
        draw.text((85, y - 9), f"{tick}%", font=fonts["small"], fill=PALETTE["gray"])
    for i, group in enumerate(groups):
        x = left + i * x_gap
        draw.line((x, top, x, bottom), fill=(220, 225, 232), width=1)
        draw_multiline_center(draw, group, x, bottom + 28, fonts["axis"], PALETTE["ink"], max_width=190)
    for si, (name, vals) in enumerate(series.items()):
        points = [(left + i * x_gap, value_to_y(value, 0, 100, top, bottom - top)) for i, value in enumerate(vals)]
        for p0, p1 in zip(points, points[1:]):
            draw.line((*p0, *p1), fill=colors[si], width=4)
        for i, (x, y) in enumerate(points):
            draw.ellipse((x - 8, y - 8, x + 8, y + 8), fill=colors[si], outline=PALETTE["white"], width=2)
            if i in {0, len(points) - 1} or name == "Sequence ASR":
                draw.text((x + 10, y - 10), f"{vals[i]:.1f}%", font=fonts["small"], fill=PALETTE["ink"])
    legend_x, legend_y = left, height - 62
    for si, name in enumerate(series):
        draw.line((legend_x, legend_y + 8, legend_x + 28, legend_y + 8), fill=colors[si], width=4)
        draw.ellipse((legend_x + 9, legend_y - 1, legend_x + 19, legend_y + 17), fill=colors[si])
        draw.text((legend_x + 38, legend_y - 2), name, font=fonts["axis"], fill=PALETTE["ink"])
        legend_x += 230
    image.save(path)


def draw_grouped_bars(
    path: Path,
    *,
    title: str,
    groups: list[str],
    series: dict[str, list[float]],
    y_label: str,
    diverging: bool,
) -> None:
    width, height = 1400, 820
    margin = dict(left=120, right=40, top=90, bottom=150)
    image = Image.new("RGB", (width, height), PALETTE["white"])
    draw = ImageDraw.Draw(image)
    fonts = chart_fonts()
    draw.text((margin["left"], 28), title, font=fonts["title"], fill=PALETTE["navy"])
    all_values = [value for values in series.values() for value in values]
    if diverging:
        min_v = min(min(all_values), -10)
        max_v = max(max(all_values), 10)
    else:
        min_v = 0
        max_v = max(max(all_values), 100)
    plot_w = width - margin["left"] - margin["right"]
    plot_h = height - margin["top"] - margin["bottom"]
    zero_y = value_to_y(0, min_v, max_v, margin["top"], plot_h)
    draw.line((margin["left"], zero_y, margin["left"] + plot_w, zero_y), fill=(180, 180, 180), width=2)
    for tick in nice_ticks(min_v, max_v, 5):
        y = value_to_y(tick, min_v, max_v, margin["top"], plot_h)
        draw.line((margin["left"] - 5, y, margin["left"] + plot_w, y), fill=(226, 230, 235), width=1)
        draw.text((20, y - 8), f"{tick:g}", font=fonts["axis"], fill=PALETTE["gray"])
    group_w = plot_w / len(groups)
    bar_w = min(34, group_w / (len(series) + 1))
    colors = [PALETTE["red"], PALETTE["gold"], PALETTE["purple"], PALETTE["green"], PALETTE["teal"]]
    for gi, group in enumerate(groups):
        center = margin["left"] + group_w * gi + group_w / 2
        start = center - (len(series) * bar_w) / 2
        for si, (name, values) in enumerate(series.items()):
            value = values[gi]
            x0 = start + si * bar_w
            x1 = x0 + bar_w * 0.75
            y = value_to_y(value, min_v, max_v, margin["top"], plot_h)
            draw.rectangle((x0, min(y, zero_y), x1, max(y, zero_y)), fill=colors[si % len(colors)])
        draw_multiline_center(draw, group, center, height - 122, fonts["axis"], PALETTE["ink"], max_width=180)
    draw.text((margin["left"], height - 38), y_label, font=fonts["axis"], fill=PALETTE["gray"])
    legend_x = margin["left"]
    legend_y = height - 75
    for si, name in enumerate(series):
        draw.rectangle((legend_x, legend_y, legend_x + 14, legend_y + 14), fill=colors[si % len(colors)])
        draw.text((legend_x + 20, legend_y - 2), name, font=fonts["axis"], fill=PALETTE["ink"])
        legend_x += 190
    image.save(path)


def draw_stacked_bars(
    path: Path,
    *,
    title: str,
    groups: list[str],
    stacks: list[str],
    values: list[list[int]],
    colors: list[tuple[int, int, int]],
) -> None:
    width, height = 1400, 760
    image = Image.new("RGB", (width, height), PALETTE["white"])
    draw = ImageDraw.Draw(image)
    fonts = chart_fonts()
    draw.text((70, 28), title, font=fonts["title"], fill=PALETTE["navy"])
    left, top, bar_w, bar_h, gap = 330, 110, 860, 54, 72
    max_total = max(sum(row) for row in values) or 1
    for i, group in enumerate(groups):
        y = top + i * gap
        draw.text((55, y + 14), group, font=fonts["axis"], fill=PALETTE["ink"])
        x = left
        total = sum(values[i])
        for j, value in enumerate(values[i]):
            seg = bar_w * value / max_total
            if seg > 0:
                draw.rectangle((x, y, x + seg, y + bar_h), fill=colors[j])
                if seg > 58:
                    draw.text((x + 6, y + 18), str(value), font=fonts["small"], fill=PALETTE["white"])
                x += seg
        draw.text((left + bar_w + 18, y + 17), f"n={total}", font=fonts["axis"], fill=PALETTE["gray"])
    legend_x, legend_y = left, height - 110
    for stack, color in zip(stacks, colors):
        if legend_x > width - 150:
            legend_x = left
            legend_y += 30
        draw.rectangle((legend_x, legend_y, legend_x + 16, legend_y + 16), fill=color)
        draw.text((legend_x + 22, legend_y - 2), stack, font=fonts["axis"], fill=PALETTE["ink"])
        legend_x += 130
    image.save(path)


def draw_horizontal_bars(
    path: Path,
    *,
    title: str,
    items: list[tuple[str, int]],
    color: tuple[int, int, int],
) -> None:
    width, height = 1400, 720
    image = Image.new("RGB", (width, height), PALETTE["white"])
    draw = ImageDraw.Draw(image)
    fonts = chart_fonts()
    draw.text((70, 28), title, font=fonts["title"], fill=PALETTE["navy"])
    left, top, bar_w, gap = 520, 105, 700, 80
    max_value = max(v for _, v in items) or 1
    for i, (label, value) in enumerate(items):
        y = top + i * gap
        draw_multiline_right(draw, label, left - 18, y + 8, fonts["axis"], PALETTE["ink"], max_width=430)
        seg = bar_w * value / max_value
        draw.rectangle((left, y, left + seg, y + 32), fill=color)
        draw.text((left + seg + 14, y + 8), str(value), font=fonts["axis"], fill=PALETTE["ink"])
    image.save(path)


def chart_fonts() -> dict[str, ImageFont.FreeTypeFont | ImageFont.ImageFont]:
    candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
    ]
    font_path = next((p for p in candidates if p.exists()), None)
    if font_path:
        return {
            "title": ImageFont.truetype(str(font_path), 30),
            "axis": ImageFont.truetype(str(font_path), 20),
            "small": ImageFont.truetype(str(font_path), 16),
        }
    default = ImageFont.load_default()
    return {"title": default, "axis": default, "small": default}


def value_to_y(value: float, min_v: float, max_v: float, top: int, plot_h: int) -> float:
    return top + (max_v - value) / (max_v - min_v) * plot_h


def value_to_x(value: float, min_v: float, max_v: float, left: int, plot_w: int) -> float:
    return left + (value - min_v) / (max_v - min_v) * plot_w


def nice_ticks(min_v: float, max_v: float, count: int) -> list[float]:
    if min_v == max_v:
        return [min_v]
    step = (max_v - min_v) / count
    magnitude = 10 ** math.floor(math.log10(abs(step)))
    nice_step = round(step / magnitude) * magnitude
    start = math.floor(min_v / nice_step) * nice_step
    ticks = []
    value = start
    while value <= max_v + nice_step:
        if min_v <= value <= max_v:
            ticks.append(round(value, 2))
        value += nice_step
    return ticks


def draw_multiline_center(draw: ImageDraw.ImageDraw, text: str, x: float, y: float, font: Any, fill: Any, max_width: int) -> None:
    lines = wrap_text(draw, text, font, max_width)
    for i, line in enumerate(lines):
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x - (bbox[2] - bbox[0]) / 2, y + i * 22), line, font=font, fill=fill)


def draw_multiline_right(draw: ImageDraw.ImageDraw, text: str, x: float, y: float, font: Any, fill: Any, max_width: int) -> None:
    lines = wrap_text(draw, text, font, max_width)
    for i, line in enumerate(lines):
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x - (bbox[2] - bbox[0]), y + i * 22), line, font=font, fill=fill)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: Any, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        bbox = draw.textbbox((0, 0), candidate, font=font)
        if bbox[2] - bbox[0] <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def add_picture(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(path), width=Inches(6.3))
    shape._inline.docPr.set("title", caption.split(".", 1)[0])
    shape._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph(caption, style="Figure Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    p.add_run(text)
    mark_header_row(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[Any]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        shade_cell(hdr[i], "F2F4F7")
        set_cell_margins(hdr[i])
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9.5)
    mark_header_row(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def set_table_widths(table: Any, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell: Any, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def mark_header_row(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_bottom_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D7DBE2")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def find_row(rows: list[dict[str, Any]], **criteria: Any) -> dict[str, Any]:
    for row in rows:
        if all(row.get(key) == value for key, value in criteria.items()):
            return row
    raise KeyError(criteria)


def metric_value(rows: list[dict[str, Any]], setting_id: str, metric: str) -> str:
    row = find_row(rows, setting_id=setting_id, metric=metric)
    return f"{float(row['mean_value']):.2f}%"


def metric_delta(rows: list[dict[str, Any]], setting_id: str, metric: str) -> str:
    row = find_row(rows, setting_id=setting_id, metric=metric)
    value = float(row["delta_vs_full"])
    return f"{value:+.2f} pp"


def top_blocked(rows: list[dict[str, Any]], setting_id: str, source_group: str) -> str:
    subset = [
        row
        for row in rows
        if row["setting_id"] == setting_id and row["source_group"] == source_group and row["blocked_at"] != "NONE"
    ]
    subset.sort(key=lambda row: int(float(row["turns"])), reverse=True)
    return ", ".join(f"{row['blocked_at']}={int(float(row['turns']))}" for row in subset[:3]) or "No dominant block"


def migration_turns(
    rows: list[dict[str, Any]],
    left: str,
    right: str,
    source_group: str,
    left_decision: str,
    right_decision: str,
) -> int:
    return sum(
        int(float(row["turns"]))
        for row in rows
        if row["left_setting_id"] == left
        and row["right_setting_id"] == right
        and row["source_group"] == source_group
        and row["left_decision"] == left_decision
        and row["right_decision"] == right_decision
    )


def migration_interpretation(left: str, right: str, source: str, left_decision: str, right_decision: str) -> str:
    if source == "ST-RBAC" and left_decision == "DENY" and right_decision == "ALLOW":
        return "Authorization evidence disappears; RBAC attacks move from blocked to served."
    if source == "MT-MAL" and left_decision == "DENY" and right_decision == "ALLOW":
        return "Multi-turn malicious turns become more permissive without downstream authorization gates."
    if source == "ST-PI" and left_decision == "ALLOW" and right_decision == "DENY":
        return "Restoring the M1/M2 combination pushes prompt-injection traffic back into denial."
    if source == "ST-RBAC" and left_decision == "ALLOW" and right_decision == "DENY":
        return "The comparison shows the removed authorization block was carrying the RBAC refusal load."
    return "Same sample_id follows a different decision path under the ablation."


def representation_change(setting_id: str) -> str:
    if setting_id == "trustedsql_minus_m1":
        return "The request skips the prompt-integrity artifact, so explicit bypass language must be discovered later by M2/M4/M5 or may reach SQL."
    if setting_id == "trustedsql_minus_m2":
        return "The runtime loses conversation-risk intent, scope, target-relation, and security-transition evidence; later policy modules still operate."
    if setting_id == "trustedsql_minus_m3_m4_m5":
        return "The pipeline no longer constructs or verifies the resource contract; SQL generation receives no deterministic table/column or row-scope proof boundary."
    if setting_id == "trustedsql_minus_m7":
        return "The system keeps upstream policy/risk gates but loses the final SQL AST conformance checkpoint before execution."
    return ""


def case_module_evidence_rows(setting: str, modules: list[dict[str, Any]]) -> list[list[str]]:
    rows: list[list[str]] = []
    for module in modules:
        module_id = module.get("module_id", "")
        evidence = ""
        if module_id == "M1":
            verdict = module.get("llm_verdict")
            reason = module.get("llm_reason")
            hits = module.get("heuristic_hits") or []
            evidence = f"verdict={verdict}; heuristic_hits={hits}; reason={reason}"
        elif module_id == "M2":
            signals = module.get("security_signals") or []
            evidence = (
                f"intent={module.get('primary_intent')}; scope={module.get('scope')}; "
                f"target={module.get('target_relation')}; transition={module.get('security_transition')}; signals={signals}"
            )
        elif module_id == "M3":
            evidence = (
                f"policy_refs={module.get('policy_refs')}; resources={module.get('requested_resources')}; "
                f"scope_type={module.get('scope_type')}; target_table={module.get('target_resource_table')}"
            )
        elif module_id == "M4":
            evidence = (
                f"reason={module.get('reason_code')}; violations={module.get('violations')}; "
                f"count={module.get('violations_count')}"
            )
        elif module_id == "M5":
            evidence = f"reason={module.get('reason_code')}; proof_status={module.get('proof_status')}"
        elif module_id == "M6":
            evidence = f"raw_sql={module.get('raw_sql')}"
        elif module_id == "M7":
            evidence = f"final_sql={module.get('final_sql')}"
        elif module_id == "X1":
            evidence = f"row_count={module.get('row_count')}; error={module.get('error')}"
        else:
            evidence = f"stage={module.get('stage')}; decision={module.get('decision')}"
        rows.append([setting, f"{module_id}:{module.get('decision')}", evidence])
    return rows


def fmt_pct(value: Any) -> str:
    return f"{float(value):.2f}%"


if __name__ == "__main__":
    raise SystemExit(main())
